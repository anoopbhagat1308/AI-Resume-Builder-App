

import streamlit as st
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Resume Builder")
st.title("📄 AI Resume Builder")

# -----------------------------
# User Inputs
# -----------------------------
name = st.text_input("Name")
email = st.text_input("Email")
phone = st.text_input("Phone")

summary = st.text_area("Professional Summary")

skills = st.text_area("Skills (comma separated)")

education = st.text_input("Education")

experience = st.text_area("Experience")

projects = st.text_area("Projects")

# -----------------------------
# Function to Create Resume
# -----------------------------
def create_resume():
    doc = Document()
    doc.add_heading(name.title(), level=1)
    doc.add_paragraph(f"Email : {email}")
    doc.add_paragraph(f"Phone : {phone}")
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(summary.title())
    doc.add_heading("Skills", level=2)

    for skill in skills.split(","):
        if skill.strip():
            doc.add_paragraph(skill.strip().title(), style="List Bullet")

    doc.add_heading("Education", level=2)
    doc.add_paragraph(education.title())

    doc.add_heading("Experience", level=2)
    doc.add_paragraph(experience.title())

    doc.add_heading("Projects", level=2)
    doc.add_paragraph(projects.title())

    file = BytesIO()
    doc.save(file)
    file.seek(0)

    return file

# -----------------------------
# Generate Resume
# -----------------------------
if st.button("Generate Resume"):
    resume = create_resume()
    st.success("Resume Generated Successfully!")
    st.download_button(
        "⬇ Download Resume",
        resume,
        file_name="Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )